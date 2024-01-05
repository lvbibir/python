#!/usr/bin/env python

# -*- coding: utf-8 -*-
"""
@File    :   main.py
@Time    :   2023/03/20 17:24:22
@Author  :   Levi Liu
@Version :   1.0
@Site    :   https://www.lvbibir.cn
@Desc    :   None
"""
import os

# from docx import Document
import docx

PWD = os.path.dirname(os.path.abspath(__file__))
os.chdir(PWD)
doc = docx.Document("test.docx")


def check_tables():
    tables = doc.tables
    for table in tables:
        print(f"表格行数: {len(table.rows)}")
        print(f"表格列数: {len(table.columns)}")
        print(f"表格单元格数量: {len(table._cells)}")
        for i in range(0, len(table.rows)):
            for j in range(0, len(table.columns)):
                print(table.cell(i, j).text, end="\t")
            print()


def add_tables():
    doc.add_table(3, 4)
    doc.save("test.docx")


def main():
    add_tables()
    check_tables()


if __name__ == "__main__":
    main()
